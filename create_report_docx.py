# -*- coding: utf-8 -*-
"""
Script tạo Báo cáo Chi tiết Chức năng Dự án Hóa Đơn / Báo Giá Mạnh Hùng (.docx)
"""

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generate_doc():
    doc = docx.Document()

    # Cấu hình lề trang A4 (2cm = 1134 dxa)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Styling cơ bản
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # ----- TIÊU ĐỀ BÁO CÁO -----
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("BÁO CÁO TỔNG QUAN VÀ CHI TIẾT TÍNH NĂNG DỰ ÁN\n")
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    subtitle_run = title_p.add_run("ỨNG DỤNG QUẢN LÝ & LẬP HÓA ĐƠN / BÁO GIÁ BÁN HÀNG PWA\n(Thế Giới Cửa Mạnh Hùng - Version 2.5 Enterprise SaaS 2025)")
    subtitle_run.font.name = 'Times New Roman'
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.bold = True
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    doc.add_paragraph() # Spacer

    # ----- BẢNG THÔNG TIN TỔNG QUAN -----
    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("Tên dự án:", "Ứng dụng Lập Hóa Đơn & Bảng Báo Giá Bán Hàng PWA"),
        ("Đơn vị sở hữu:", "Thế Giới Cửa Mạnh Hùng (SĐT: 0976.088.080 - 0916.557.888)"),
        ("Địa chỉ Web Live:", "https://rycodeapp.github.io/hoa-don-manh-hung/"),
        ("Mã nguồn GitHub:", "https://github.com/Rycodeapp/hoa-don-manh-hung"),
        ("Công nghệ cốt lõi:", "HTML5, Vanilla JS (ES6 Modular), TailwindCSS 2025, PWA Offline Service Worker")
    ]
    for row_idx, (label, val) in enumerate(info_data):
        row = info_table.rows[row_idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        
        cell_lbl.paragraphs[0].add_run(label).bold = True
        cell_lbl.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        cell_val.paragraphs[0].add_run(val)
        
        set_cell_background(cell_lbl, "F1F5F9")
        set_cell_background(cell_val, "FFFFFF")
        set_cell_margins(cell_lbl, top=120, bottom=120, left=150, right=150)
        set_cell_margins(cell_val, top=120, bottom=120, left=150, right=150)

    doc.add_paragraph()

    def add_section_header(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
        r = h.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        return h

    def add_feature_box(num, title, desc_lines):
        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_before = Pt(10)
        p_title.paragraph_format.space_after = Pt(2)
        r_num = p_title.add_run(f"Tính năng {num}: ")
        r_num.font.bold = True
        r_num.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        
        r_t = p_title.add_run(title)
        r_t.font.bold = True
        r_t.font.size = Pt(12.5)
        r_t.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

        for line in desc_lines:
            p_desc = doc.add_paragraph()
            p_desc.paragraph_format.left_indent = Inches(0.25)
            p_desc.paragraph_format.space_after = Pt(2)
            r_bullet = p_desc.add_run("• ")
            r_bullet.bold = True
            r_bullet.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
            
            parts = line.split(":", 1)
            if len(parts) == 2:
                r_key = p_desc.add_run(parts[0] + ":")
                r_key.bold = True
                p_desc.add_run(parts[1])
            else:
                p_desc.add_run(line)

    add_section_header("I. DANH SÁCH CHI TIẾT 10 NHÓM TÍNH NĂNG HIỆN CÓ")

    add_feature_box(1, "Lập Hóa Đơn & Bảng Báo Giá Tùy Chỉnh Linh Hoạt", [
        "Thay đổi Tiêu đề Chứng từ thủ công: Cho phép người dùng chuyển đổi tự do tiêu đề giữa HÓA ĐƠN, BÁO GIÁ, BẢNG BÁO GIÁ, PHIẾU GIAO HÀNG, HÓA ĐƠN BÁN HÀNG...",
        "Tự động sinh Mã chứng từ chuẩn: Hệ thống tự động tạo mã ngẫu nhiên dạng HD[ngày][tháng]-[mã số] (Ví dụ: HD2307-482) giúp tránh trùng lặp.",
        "Đồng bộ ngày lập chứng từ: Tự động ghi nhận ngày lập chứng từ theo thời gian thực và cho phép chỉnh sửa ngày tháng."
    ])

    add_feature_box(2, "Quản Lý Thông Tin Khách Hàng & Điều Khoản Bán Hàng", [
        "Quản lý Đơn vị bán hàng & SĐT: Lưu giữ tên đơn vị 'Thế giới Cửa Mạnh Hùng' và hotline liên hệ (0976.088.080 - 0916.557.888).",
        "Quản lý Khách hàng & Công trình: Cho phép nhập Họ tên người mua, Số điện thoại di động, Địa chỉ giao hàng/công trình và Kho xuất hàng.",
        "Ghi chú chung & Điều khoản: Khối nhập liệu dành riêng cho các quy định về Bảo hành (motor, thân cửa), Thuế VAT, Phí vận chuyển, Điều khoản thanh toán..."
    ])

    add_feature_box(3, "Nhập Liệu & Tính Toán Kích Thước Cửa Tự Động", [
        "Tính toán Diện tích Cửa (Ngang x Cao): Tự động nhân kích thước Ngang (m) x Cao (m) ra tổng diện tích m² chính xác đến 2 chữ số thập phân.",
        "Tính năng Nhập 'Số bộ cửa cùng kích thước': Trường hợp có 3 bộ cửa cùng số đo (3.5m x 3.2m), chỉ cần nhập Số bộ = 3, ứng dụng tự động nhân Tổng diện tích = 3.5 x 3.2 x 3 = 33.60 m².",
        "Hỗ trợ Đa dạng Đơn vị tính: Linh hoạt lựa chọn m², Bộ, Cái, Mét, kg phù hợp cho cả cửa cuốn, cửa kéo, phụ kiện và chi phí nhân công.",
        "Danh mục Sản phẩm gợi ý sẵn (Datalist): Tích hợp danh sách mẫu các loại cửa phổ biến (Cửa cuốn Đức khe thoáng, Cửa cuốn Đài Loan lá 7dem/8dem, Cửa kéo Đài Loan, Motor 400kg/600kg, Bình lưu điện, Remote...)."
    ])

    add_feature_box(4, "Chuyển Đổi Số Tiền Thành Chữ Tiếng Việt (Vietnamese Money Words)", [
        "Tự động đọc số tiền thành chữ: Chuyển đổi chính xác 100% con số tổng tiền hàng (Ví dụ: 145.752.000 đ -> 'Một trăm bốn mươi lăm triệu bảy trăm năm mươi hai ngàn đồng.').",
        "Đảm bảo chuẩn mực pháp lý: Giúp hóa đơn/báo giá minh bạch, đúng quy định tài chính kế toán tại Việt Nam."
    ])

    add_feature_box(5, "Cơ Chế Tự Động Co Giãn Cỡ Chữ Thông Minh (Auto Font Scaling)", [
        "Phân tích độ dài mô tả quy cách: Tự động đo độ dài chuỗi ký tự của từng món hàng để điều chỉnh cỡ font (14px -> 12px -> 11px -> 10px).",
        "Cam kết hiển thị 100% trọn vẹn chữ: Những mô tả kỹ thuật phức tạp (dòng cửa nhôm Xingfa hệ 55, kính uốn vòm, kính hộp, phụ kiện JANUS...) được hiển thị đầy đủ từng chữ, loại bỏ hoàn toàn lỗi bị ẩn chữ hay dấu ba chấm (...).",
        "Mật độ in tối ưu (Dense Sheet): Khi danh sách có từ 9 sản phẩm trở lên, ứng dụng tự kích hoạt chế độ dồn dòng giúp tờ hóa đơn luôn nằm gọn trên 1 trang A4."
    ])

    add_feature_box(6, "Xem Trước Hóa Đơn Paper Preview & Xuất PDF Đa Nền Tảng", [
        "Mô phỏng tờ Hóa đơn thực tế: Hiển thị giao diện xem trước mẫu giấy A4 font chuẩn Times New Roman sắc nét.",
        "Tự động gán Tên File PDF thông minh: Khi bấm Lưu PDF trên máy tính hay điện thoại, tên file tự động điền dạng [Mã Số] - [Tên Khách Hàng] - [Tiêu Đề Chứng Từ].pdf (Ví dụ: HD2307-482 - Nguyen Van A - BAAO GIA.pdf).",
        "Chuẩn hóa tên file an toàn (Sanitize Encoding): Tự động lọc sạch ký tự tiếng Việt có dấu phức tạp để tránh bị lỗi đặt tên file trên iOS Safari và Android Chrome."
    ])

    add_feature_box(7, "Giao Diện Đạt Chuẩn Enterprise SaaS 2025", [
        "Thiết kế Responsive Mobile-First: Tương thích hoàn hảo trên 11 cấp độ màn hình từ 360px (Samsung S21 FE), 393px, 768px, 1024px đến 1920px Desktop. Cam kết Zero Horizontal Scroll.",
        "Sticky Action Bar Cân Đối (45% | 10% | 45%): Thanh công cụ cố định dưới đáy với bố cục CSS Grid 3 cột trên Desktop và 2 hàng nút 50/50 trên Mobile.",
        "Nút Thêm Sản Phẩm Thông Minh kép: Đặt nút 'THÊM CỬA / SẢN PHẨM MỚI' ở cả đầu trang và cuối danh sách sản phẩm giúp thao tác cực nhanh.",
        "Thẻ Card Bo Góc Premium: Bo góc 16px, shadow nhẹ, ô nhập 52px phản hồi mịn 200ms."
    ])

    add_feature_box(8, "Khả Năng Hoạt Động Offline 100% Qua PWA Service Worker", [
        "Tích hợp PWA Manifest & App Icon: Cho phép người dùng bấm 'Cài đặt App' để đưa ứng dụng ra màn hình chính điện thoại như một App native.",
        "Service Worker Cache Network-First (v3): Tự động lưu toàn bộ file JS/CSS/HTML vào bộ nhớ cache, cho phép mở app và lập hóa đơn mượt mà ngay cả khi không có kết nối mạng Internet."
    ])

    add_feature_box(9, "Quản Lý Trạng Thái & Hệ Thống Hoàn Tác (State & Undo System)", [
        "Tính năng Hoàn tác (Undo): Cho phép khôi phục lại dữ liệu trước đó khi lỡ tay xóa dòng hoặc sửa nhầm thông tin.",
        "Tự động lưu tạm dữ liệu (LocalStorage): Mọi dữ liệu đang nhập dở sẽ tự động được bảo toàn, không lo bị mất khi lỡ tay đóng trình duyệt.",
        "Tạo hóa đơn mới an toàn: Nút 'Tạo hóa đơn mới' dọn dẹp sạch form có thông báo xác nhận an toàn."
    ])

    add_feature_box(10, "Tối Ưu In Ấn Chuẩn Mẫu Giấy A4 (@media print)", [
        "Định dạng lề in 10mm: Tối ưu lề in trên dưới trái phải, không làm mất chữ viền tờ giấy.",
        "Khối chữ ký gắn liền (Avoid Page Break): Khối tổng tiền bằng chữ, ghi chú và 3 cột chữ ký (Người nhận hàng, Đã nhận đủ tiền, Người viết hóa đơn) được khóa liền khối, không bao giờ bị tràn sang trang thứ 2 một cách rời rạc."
    ])

    doc.add_paragraph()

    add_section_header("II. KIẾN TRÚC KỸ THUẬT VÀ TÍNH ĐẢM BẢO CHIỀU SÂU MÃ NGUỒN")

    tech_p = doc.add_paragraph()
    tech_p.paragraph_format.left_indent = Inches(0.2)
    tech_p.paragraph_format.space_after = Pt(4)
    tech_p.add_run("Dự án được xây dựng theo mô hình kiến trúc ").font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    r_arch = tech_p.add_run("Modular Architecture (ES6 IIFE Pattern)")
    r_arch.bold = True
    r_arch.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    tech_p.add_run(" phân tách sạch sẽ giữa Logic trạng thái và Giao diện UI:")

    # Table cấu trúc thư mục
    struct_table = doc.add_table(rows=6, cols=2)
    struct_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    struct_data = [
        ("assets/js/state.js", "Quản lý trạng thái tập trung (Single Source of Truth), quản lý Undo stack & Subscriber pattern."),
        ("assets/js/calculator.js", "Module tính toán diện tích cửa, số bộ nhân tổng m² và cộng tổng hóa đơn."),
        ("assets/js/vietnameseMoney.js", "Module thuật toán đổi chữ số thành văn bản tiếng Việt."),
        ("assets/components/productRow.js & productTable.js", "Component dựng bảng sản phẩm, ô nhập số bộ, kích thước & tính tiền thời gian thực."),
        ("assets/components/preview.js & printer.js", "Component xem trước tờ hóa đơn/báo giá, tự co font và xử lý in PDF."),
        ("service-worker.js & manifest.json", "Cấu hình PWA Offline-First, quản lý bộ nhớ đệm Cache v3.")
    ]

    for row_idx, (fpath, fdesc) in enumerate(struct_data):
        row = struct_table.rows[row_idx]
        c1, c2 = row.cells[0], row.cells[1]
        
        c1.paragraphs[0].add_run(fpath).bold = True
        c1.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        c2.paragraphs[0].add_run(fdesc)
        
        bg = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        set_cell_background(c1, bg)
        set_cell_background(c2, bg)
        set_cell_margins(c1, top=100, bottom=100, left=120, right=120)
        set_cell_margins(c2, top=100, bottom=100, left=120, right=120)

    doc.add_paragraph()

    add_section_header("III. KẾT LUẬN VÀ HƯỚNG DẪN KHỞI CHẠY")
    
    final_p = doc.add_paragraph()
    final_p.paragraph_format.left_indent = Inches(0.2)
    final_p.add_run("Dự án ").font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    r_f1 = final_p.add_run("Thế Giới Cửa Mạnh Hùng - Hóa Đơn & Báo Giá Bán Hàng")
    r_f1.bold = True
    final_p.add_run(" là một sản phẩm hoàn chỉnh, vừa đạt độ hoàn thiện cao về mặt thị giác (Enterprise SaaS 2025 Standard), vừa xử lý triệt để các bài toán thực tế của ngành cửa (số bộ cửa, kích thước Ngang x Cao, in ấn A4 không mất chữ, lưu PDF tên khách hàng trên điện thoại).")

    output_path = r"c:\Users\MANH\Desktop\Study\Hóa Đơn\Bao_Cao_Tinh_Nang_Hoa_Don_Manh_Hung.docx"
    doc.save(output_path)

if __name__ == "__main__":
    generate_doc()
